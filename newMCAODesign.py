#!/usr/bin/env python3
# newMCAODesign.py

# generates a word document containing animal inclusion/exclusion
# criteria and a description of surgical and post-op procedure

import docx, datetime

options = {
    'Species':['wt', 'TRPm2-KO', 'Nestin', 'TRPm2-FL',
               'CAMKII-CRE', 'FL//CRE', 'BDNF', 'ERA'],
    'Sex':['males', 'females', 'males and females'],
    'Age':{'P10':['10-11 days', '5-9g', '5%', '3%', '.17', 0],
           'P21-25':['21-25 days', '10-20g', '4.5%', '2.5%', '.17', .5],
           'Adult':['8-12 weeks', '20-30g', '4%', '2%', '.21', 1]},
    'Procedure': ['45', '60', '90']
    }

def getSpecies():
    species = input('Input a strain from the list %s: \n' % options['Species'])
    if species not in options['Species']:
        print('Must select a strain from the list...\n')
        getSpecies()
    else:
        return species

def getSex():
    sex = input('Input a sex from the list %s: \n' % options['Sex'])
    if sex not in options['Sex']:
        print('Must select a sex from the list...\n')
        getSex()
    else:
        return sex

def getAge():
    age = input('Input an age from the list: %s\n' % options['Age'].keys())
    if age not in options['Age']:
        print('Must select an age from the list...\n')
        getAge()
    else:
        return age

def getPro():
    pro = input('Input a stroke time from the list: %s\n' % options['Procedure'])
    if pro not in options['Procedure']:
        print('Must select a stroke time from the list...\n')
        getPro()
    else:
        return pro

doc = docx.Document()

code = input('What is the project code?\n')

n = input('How many animals should be included?\n')

species = getSpecies()

sex = getSex()

age = getAge()

pro = getPro()

codeString = '%s Study Design' % code

today = datetime.date.today()

dayString = today.strftime("%d/%m/%Y")

codeLine = doc.add_paragraph(codeString)
codeLine.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

name = doc.add_paragraph('Ben Wassermann')
name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

date = doc.add_paragraph(dayString)
date.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT

line1 = '\nStudy will include %s animals.' % n
para1 = doc.add_paragraph(line1)

line2 = 'Mice must be %s %s, %s of age, between %s.' % (species, sex, options['Age'][age][0], options['Age'][age][1])
para2 = doc.add_paragraph(line2)

para3 = doc.add_paragraph('Surgical Procedure:')
para3.runs[0].underline = True

line4 = 'Isoflurane induction at %s, hold in nose cone at %s' % (options['Age'][age][2], options['Age'][age][3])
para4 = doc.add_paragraph(line4, style='ListNumber')

para5 = doc.add_paragraph('LDF incision opened between right ear and eye, cauterized.', style='ListNumber')

line6 = '%s-minute MCAo with a %s filament' % (pro, options['Age'][age][4])
para6 = doc.add_paragraph(line6, style='ListNumber')

line7 = '2-days post-op care consisting of moist chow and %s mL IP saline injections' % options['Age'][age][5]
para7 = doc.add_paragraph(line7, style='ListNumber')

para8 = doc.add_paragraph('Exclusion Criteria:')
para8.runs[0].underline = True

line9 = doc.add_paragraph('< 90% LDF reduction at MCAo', style='ListBullet')

line10 = doc.add_paragraph('< 50% LDF recovery at reperfusion', style='ListBullet')

line11 = doc.add_paragraph('> 25% weight loss post-op', style='ListBullet')

line12 = doc.add_paragraph('Significant behavior signals', style='ListBullet')

line13 = doc.add_paragraph('Evidence of brain hemorrhage during harvest', style='ListBullet')


fileName = '%sStudyDesign.docx' % code

doc.save(fileName)


